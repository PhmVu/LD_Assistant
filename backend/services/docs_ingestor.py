from __future__ import annotations

from datetime import datetime
import base64
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

from core.config import settings
from services.docs_index import DocsIndex
from services.ld_ai.doc_summarizer import extract_rd_pack, summarize_image, summarize_text
from services.ld_ai.knowledge_retriever import clear_knowledge_cache
from services.ld_ai.synaptic_vortex import get_global_vortex

DOCX_EXTS = {".docx"}
XLSX_EXTS = {".xlsx", ".xlsm", ".xltx", ".xltm"}
TEXT_EXTS = {".txt", ".md"}
IMAGE_EXTS = {".png", ".jpg", ".jpeg", ".gif", ".webp"}


def _safe_name(filename: str) -> str:
    return filename.replace("/", "_").replace("\\", "_")


def _truncate(text: str, max_len: int = 500) -> str:
    clean = (text or "").strip()
    return clean[:max_len] + ("..." if len(clean) > max_len else "")


def _infer_marking_type(text: str) -> str:
    t = (text or "").lower()
    if "nét đứt" in t or "dashed" in t:
        return "dashed"
    if "vạch liền" in t or "solid" in t:
        return "solid"
    if "lề" in t or "edge" in t or "curb" in t:
        return "edge"
    if "mũi tên" in t or "arrow" in t:
        return "arrow"
    if "zebra" in t or "crosswalk" in t:
        return "crosswalk"
    if "vạch dừng" in t or "stop line" in t:
        return "stop_line"
    return "general"


def _extract_keywords(text: str, limit: int = 18) -> List[str]:
    tokens = [t.strip(".,;:()[]{}\n\t\r") for t in (text or "").lower().split()]
    filtered = [t for t in tokens if len(t) > 2]
    seen: List[str] = []
    for token in filtered:
        if token not in seen:
            seen.append(token)
        if len(seen) >= limit:
            break
    return seen


def _save_asset(doc_name: str, idx: int, blob: bytes, ext: str) -> Path:
    assets_dir = settings.DOCS_ASSETS_DIR
    assets_dir.mkdir(parents=True, exist_ok=True)
    safe_doc = Path(doc_name).stem.replace(" ", "_")
    asset_name = f"{safe_doc}_img_{idx}{ext}"
    path = assets_dir / asset_name
    path.write_bytes(blob)
    return path


def _extract_docx(path: Path) -> Tuple[str, List[Tuple[bytes, str]]]:
    try:
        from docx import Document
    except Exception:
        return "", []

    doc = Document(str(path))
    parts: List[str] = []
    for para in doc.paragraphs:
        if para.text and para.text.strip():
            parts.append(para.text.strip())
    for table in doc.tables:
        for row in table.rows:
            row_vals = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_vals:
                parts.append(" | ".join(row_vals))

    images: List[Tuple[bytes, str]] = []
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            try:
                blob = rel.target_part.blob
                content_type = rel.target_part.content_type or "image/png"
                images.append((blob, content_type))
            except Exception:
                continue

    return "\n".join(parts), images


def _extract_xlsx(path: Path) -> Tuple[str, List[Tuple[bytes, str]]]:
    try:
        import openpyxl
    except Exception:
        return "", []

    wb = openpyxl.load_workbook(str(path), data_only=True)
    parts: List[str] = []
    images: List[Tuple[bytes, str]] = []

    for ws in wb.worksheets:
        parts.append(f"[SHEET] {ws.title}")
        for row in ws.iter_rows(values_only=True):
            if not row:
                continue
            values = [str(v).strip() for v in row if v is not None and str(v).strip()]
            if values:
                parts.append("\t".join(values))

        for img in getattr(ws, "_images", []):
            try:
                blob = img._data()  # type: ignore[attr-defined]
                images.append((blob, "image/png"))
            except Exception:
                continue

    return "\n".join(parts), images


def _summarize_images(doc_name: str, images: List[Tuple[bytes, str]]) -> List[Dict[str, Any]]:
    results: List[Dict[str, Any]] = []
    for idx, (blob, content_type) in enumerate(images, start=1):
        try:
            b64 = base64.b64encode(blob).decode("utf-8")
            preview = summarize_image(b64) or ""
        except Exception:
            preview = ""
        ext = ".png" if "png" in content_type else ".jpg"
        path = _save_asset(doc_name, idx, blob, ext)
        results.append({"path": str(path), "preview": preview, "content_type": content_type})
    return results


def ingest_file(path: Path, content_type: Optional[str] = None) -> Dict[str, Any]:
    docs_index = DocsIndex()
    vortex = get_global_vortex()

    safe_name = _safe_name(path.name)
    ext = path.suffix.lower()
    text = ""
    image_items: List[Dict[str, Any]] = []

    if ext in DOCX_EXTS:
        text, images = _extract_docx(path)
        image_items = _summarize_images(safe_name, images)
    elif ext in XLSX_EXTS:
        text, images = _extract_xlsx(path)
        image_items = _summarize_images(safe_name, images)
    elif ext in TEXT_EXTS:
        text = path.read_text(encoding="utf-8", errors="ignore")
    elif ext in IMAGE_EXTS:
        blob = path.read_bytes()
        image_items = _summarize_images(safe_name, [(blob, content_type or "image/png")])

    if text:
        try:
            summary = summarize_text(text)
        except Exception:
            summary = _truncate(text)
        try:
            rd_pack = extract_rd_pack(text)
        except Exception:
            rd_pack = None
    else:
        summary = ""
        rd_pack = None
    preview = _truncate(summary or text)
    marking_type = _infer_marking_type(text)
    keywords = _extract_keywords(text)

    text_path = None
    if text:
        text_path = settings.DOCS_DIR / f"{Path(safe_name).stem}.txt"
        text_path.write_text(text, encoding="utf-8")

    doc_meta = {
        "name": safe_name,
        "size": path.stat().st_size,
        "content_type": content_type or "application/octet-stream",
        "updated": datetime.utcnow().strftime("%Y-%m-%d %H:%M:%S"),
        "path": str(path),
        "preview": preview,
        "summary": summary,
        "rd_pack": rd_pack,
        "text_path": str(text_path) if text_path else None,
        "image_previews": image_items,
        "keywords": keywords,
        "marking_type": marking_type,
    }

    docs_index.upsert_doc(doc_meta)
    clear_knowledge_cache()

    if summary:
        try:
            vortex.add(
                f"MARKING::{marking_type}",
                {"summary": summary, "source": safe_name, "rd_pack": rd_pack},
            )
            vortex.add(
                f"DOC::{safe_name}",
                {"summary": summary, "source": safe_name, "rd_pack": rd_pack},
            )
        except Exception:
            pass

    for img in image_items:
        if img.get("preview"):
            try:
                vortex.add(
                    f"DOC::{safe_name}",
                    {"summary": img.get("preview"), "source": img.get("path")},
                )
            except Exception:
                pass

    return doc_meta


def ingest_directory(folder: Path) -> List[Dict[str, Any]]:
    if not folder.exists():
        return []
    results: List[Dict[str, Any]] = []
    for path in folder.iterdir():
        if path.is_file():
            results.append(ingest_file(path))
    return results
